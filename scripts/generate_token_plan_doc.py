#!/usr/bin/env python3
"""
生成 MicroBubble Agent 项目开发状况报告 (Word 文档)
用于申请更大的 token plan 计划

数据来源: CLAUDE.md (项目上下文) + memory/ (75+ 沉淀)
生成日期: 2026-06-30
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 文档初始化
# ============================================================
doc = Document()

# 设置中文字体
def set_chinese_font(run, font_name="Microsoft YaHei", size=10.5, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    if color:
        run.font.color.rgb = color


# ============================================================
# 文档标题
# ============================================================
title = doc.add_heading('', level=0)
title_run = title.add_run('MicroBubble Agent 项目开发状况报告')
set_chinese_font(title_run, size=22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('（用于申请更大 Token Plan 计划）')
set_chinese_font(subtitle_run, size=12, bold=False, color=RGBColor(0x59, 0x59, 0x59))

# 日期
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run('报告日期：2026-06-30')
set_chinese_font(date_run, size=10, color=RGBColor(0x80, 0x80, 0x80))

doc.add_paragraph()


# ============================================================
# 辅助函数
# ============================================================
def add_h1(text):
    h = doc.add_heading('', level=1)
    h_run = h.add_run(text)
    set_chinese_font(h_run, size=18, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    return h


def add_h2(text):
    h = doc.add_heading('', level=2)
    h_run = h.add_run(text)
    set_chinese_font(h_run, size=15, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
    return h


def add_h3(text):
    h = doc.add_heading('', level=3)
    h_run = h.add_run(text)
    set_chinese_font(h_run, size=13, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    return h


def add_para(text, bold=False, size=10.5, color=None, italic=False):
    p = doc.add_paragraph()
    p_run = p.add_run(text)
    set_chinese_font(p_run, size=size, bold=bold, color=color)
    if italic:
        p_run.italic = True
    return p


def add_bullet(text, level=0, bold=False, size=10.5, color=None):
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(0.5 * (level + 1))
    p_run = p.add_run(text)
    set_chinese_font(p_run, size=size, bold=bold, color=color)
    return p


def add_table_header(table, headers):
    """添加表头"""
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        set_chinese_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # 设置背景色
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1F4E79')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, row_data, font_size=10):
    """添加表格行"""
    row = table.add_row()
    for i, cell_text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        set_chinese_font(run, size=font_size)


# ============================================================
# 第一部分：项目概览
# ============================================================
add_h1('一、项目概览')

add_h2('1.1 项目简介')
add_para('"小气" — 微纳米气泡课题组智能 Agent 系统，约 20 人研究实验室的 AI 助手。')

add_h2('1.2 技术栈')
add_bullet('后端：Python 3.11 + FastAPI + SQLAlchemy + PostgreSQL + Redis + Celery', bold=True)
add_bullet('前端：Vue 3 + Vite + Element Plus（桌面端） + NutUI 4（移动端）', bold=True)
add_bullet('AI：Claude API (Sonnet 4.5 + Haiku) + faster-whisper + pgvector + text2vec-base-chinese', bold=True)
add_bullet('部署：云服务器（Nginx + FRP + Webhook 9001）+ 本地电脑（Docker 9 services + GPU Whisper）', bold=True)
add_bullet('代码托管：GitHub (gg320324492-lgtm/microbubble-agent) + Webhook 自动部署', bold=True)

add_h2('1.3 项目规模指标')
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
add_table_header(table, ['指标', '数值'])
metrics = [
    ('总代码量', '313,000+ 行（80 万字符）'),
    ('文件数', '799+ 个'),
    ('Git commits 累计', '1,545+ 个'),
    ('开发天数', '46 天'),
    ('CLAUDE.md 沉淀', '566 KB（含 50+ 章节铁律）'),
    ('memory/ 沉淀', '75+ 条事实记忆'),
    ('scripts/ 工具脚本', '25+ 个'),
    ('docs/ 设计文档', '21+ 篇'),
    ('alembic 迁移', '40+ 个版本'),
    ('pytest 测试', '8+ 套件 / 多个用例'),
    ('Vitest 测试', '492+ 个用例'),
    ('qa-bench 题库', '910 题（含 535 题 7 维评分体系）'),
    ('Docker 服务', '9 个（app / db / redis / minio / neo4j / whisper / vision-mcp / celery-worker / celery-beat）'),
    ('@tool 装饰器工具', '34 个（任务 5 / 会议 7 / 项目 3 / 成员 2 / 知识 9 / 公式 1 / 假设 1 / 记忆 3 / 搜索 1 / 个性化 2 / 反馈 1）'),
    ('Rich Block 组件', '12 类（meeting / task_list / knowledge_ref / member / formula / hypothesis / project / transcript / chart / 2 兜底）'),
    ('CSS 主题', '6 套（3 主色 × 2 明暗）'),
    ('dark mode 覆盖', '桌面 11/15 view + 移动 25/25 view + 11/11 Rich Block'),
    ('Stylelint 错误', '0 errors（基线 + trend）'),
    ('Playwright baseline', '10+ 路由'),
]
for k, v in metrics:
    add_table_row(table, [k, v])

doc.add_paragraph()


# ============================================================
# 第二部分：核心架构
# ============================================================
add_h1('二、核心架构')

add_h2('2.1 智能体架构（方案 C 收官）')
add_para('采用单阶段流式渐进综合架构：取消 brief/detail 双层 → intent → agentic_loop → critique → done。')
add_para('6 条铁律：')
add_bullet('跨 event loop 安全（ToolContext 注入）', bold=True)
add_bullet('typing import CI 检查（避免模块加载 NameError）', bold=True)
add_bullet('SSE 事件 delta 语义显式标注 [increment] 或 [snapshot]', bold=True)
add_bullet('流式 abort 安全（TraceCollector + CancelledError partial）', bold=True)
add_bullet('LLMClient interface keyword-only（防静默走错模型）', bold=True)
add_bullet('feature flag 保留老路径（30 天回滚窗口）', bold=True)

add_h2('2.2 前端路由级双栈')
add_para('桌面端（Element Plus）和移动端（NutUI 4）同一 URL 不同组件，桌面 el-* 与移动 nut-* CSS 完全隔离。')

add_h2('2.3 知识库大脑（Knowledge Brain）')
add_para('八大模块：动态 LLM 分析 / 自动关联引擎 / RAG 问答引擎 / 自主研究引擎 / 健康监控 / 实体知识图谱 / 假设生成引擎 / 量化推理引擎。')

add_h2('2.4 限流基建 v31.2.x')
add_para('5 个独立 tier（auth 20/min / write 30/min / read 200/min / upload 10/min / sse 10/min）+ Redis ZSET 滑动窗口持久化 + X-RateLimit-Policy 头 + user_id 维度。')


# ============================================================
# 第三部分：已完成的开发内容（核心成果）
# ============================================================
doc.add_page_break()
add_h1('三、已完成的开发内容（核心成果）')

# --- 3.1 2026-06-30 当日收官 ---
add_h2('3.1 2026-06-30 当日收官（13 项任务）')

add_h3('3.1.1 qa-bench v3.0 完整收官（6 周 T6.1-T6.5）')
add_bullet('700 题题库（业务 500 + P 高级 100 + K 横切 100）', bold=False)
add_bullet('3 个 P0 检测器：stream_interrupt / tool_error_propagated / first_token_latency', bold=False)
add_bullet('7 维评分：intent/tool/content/rich/defense/perf/consistency + 一票否决', bold=False)
add_bullet('Dashboard MVP：Chart.js 4 图 + 失败题 TOP 20', bold=False)
add_bullet('GitHub Actions CI smoke 200 题 5min 80% 阈值', bold=False)
add_bullet('4 铁律（dict key 类别冲突 K vs D / veto 用 <= / 检测器加载失败 try-except / 短路流 text_delta>5 豁免）', bold=False)
add_para('W2 收官：229 手工题 + 107 DB 题 + 144 模板 = 535 题合并去重', italic=True)
add_para('W3 收官：发现 Self-RAG #009 回归 bug (MicroBubbleAgent.chat_stream 缺 model 参数)', italic=True)
add_para('W4 收官：P 高级 102 题 + 3-tier 阈值分档实施在 agentic_loop.py:615-665', italic=True)
add_para('W5 收官：save_to_kb.py 5 道防线 + scripts/auto_intake_rollback.py 7 天自动清理 + 200 题 smoke 套件', italic=True)
add_para('W6 收官：7 维雷达图（content 53% 是短板）+ ROI 100-150% + 4 部分 SOP + 6 周总览', italic=True)

add_h3('3.1.2 KB 入库监控（W6 D5）')
add_bullet('后端 GET /api/v1/knowledge/auto-intake-summary（today_intake + weekly_intake[7] + hit_rate + negative_feedback_rate + rollback_count + total_in_db=179）')
add_bullet('前端 web/src/composables/useKbMonitor.js（polling 5min Q5 实现 setInterval(fetchSummary, 300000)）')
add_bullet('ProjectStatsView 第 3 个 tab（4 metric card + 7 日趋势 CSS 柱状图 + 系统状态卡）')
add_bullet('2 铁律（lightningcss 不支持自定义 @keyframes 改用 Element Plus is-loading / 错误时保留旧数据不清空）')

add_h3('3.1.3 #009 Self-RAG 重检索 + 用户深度思考开关（4 commits 收官）')
add_bullet('Phase 0.5 gate (Haiku judge + deterministic 改写 + re-retrieve)')
add_bullet('前端 useUiStore useDeepThinking + ToolContext 2 kwarg 透传 + 7 AGENT_SELF_RAG_* flag')
add_bullet('8 铁律：tool_use 配对 / default-on-fail / 双层控制 per-request+settings / rerank_score 合并去重 / reretrieve_ 前缀 / 只对 search_info+explain_concept / user_message 透传 / dark 非 scoped 块')
add_bullet('6 踩坑教训：judge 800ms 太慢 / refined_query 比 original 差 / 合并 id 冲突 / default-on-fail 误通过 / 端点 Optional 兼容')

add_h3('3.1.4 v78 UI redesign 3-zone + EP icons + 4-attr a11y（4 commits 收官）')
add_bullet('SessionSidebar overlap 修复（flex min-width:0）')
add_bullet('actions 改 right-click/long-press 上下文菜单 + sortedSessions 置顶冒泡')
add_bullet('NavRail.vue 新左侧 nav rail + ChatViewSSE 3-zone 重构（≡ / ChatBreadcrumb / +FAB）')
add_bullet('ThinkingModeSwitch segmented（替代 🧠🧠 双 toggle 冲突）')
add_bullet('移动端 EP icons 同步 + variables.css --icon-size-* token')
add_bullet('8 条铁律（含绝对定位+hover-only 必重叠 / 双 toggle 同 emoji 必冲突 / EP icon 永远优于 emoji / a11y 4-attr 100% / sortedSessions pin-bubble / flex min-width:0 是 ellipsis 充要条件 / dark 非 scoped 块 v60-v67 / 单 + FAB 替代语义模糊 button）')

add_h3('3.1.5 voiceprint 视觉收官（5 commits + 1 合并）')
add_bullet('VoiceprintCard class 化 + VoiceTestDialog Canvas getComputedStyle + ConfidenceChart ECharts 主题色 + Vitest 阈值 8 个 + Playwright 6 主题 smoke test')
add_bullet('5 条新铁律（Canvas 不支持 var() / 主题切换必须主动重绘 / 颜色收敛 3 优先级 / per-card max 保留 + 全局阈值 class / 任务号冲突处理）')

add_h3('3.1.6 KB 数据清洁 B+C 方案')
add_bullet('B 物理删 1 条字节相同副本：5 类 FK 引用全扫（knowledge_relations/images/extractions/gaps ARRAY/rag_evaluations Text）+ 保守策略任一引用整组跳过 + JSON 文件备份 28936 字节')
add_bullet('C 前端 dedup toggle 可逆显示：dedupEnabled props + toggle-dedup emit + displayedItems computed + 49 title 重复的 144 条按 created_at 降序，ON 时按 title 分组取 id 最小')
add_bullet('8 条新铁律：FK 防御不是可选项 / knowledge_gaps.knowledge_ids 是 ARRAY 不是 FK / rag_evaluations.context ILIKE 数字边界陷阱 / 同 title 但 md5 不全同 → 整组保留 / DELETE 不可逆 → JSON 备份是底线 / dedup toggle 是"显示策略"不是"数据操作" / localStorage key 必须带项目前缀 / 持久化同步是单向的 (UI → localStorage)')

add_h3('3.1.7 180 张 [拓展-XX] 卡片 source_type 重分类')
add_bullet('NULL → auto_expansion 防御性 WHERE 过滤')
add_bullet('踩坑 1：SQLAlchemy regex 转义陷阱 → 改用 String.startswith')
add_bullet('新铁律 7：SQLAlchemy regex 转义陷阱')

add_h3('3.1.8 conftest.py test_member 隔离 bug 修复（yield-teardown + replica role）')
add_bullet('SQLAlchemy session.rollback() 不撤销 committed row')
add_bullet('fixture return 等同"无 teardown"，必须用 yield 才有 cleanup 钩子')
add_bullet('PostgreSQL FK RESTRICT 必须用 SET session_replication_role = replica 绕过')
add_bullet('pytest fixture teardown 顺序 = 反向 setup (LIFO)')
add_bullet('test_auth.py 4/8 FAIL → 8/8 PASS')

add_h3('3.1.9 #042 概念问 4 域代码强制 fan-out（commit 5522ad5a 夹带收官）')
add_bullet('app/config.py: AGENT_CROSS_DOMAIN_FANOUT_ENABLED flag')
add_bullet('app/agent/agentic_loop.py: CONCEPT_DOMAIN_TOOLS + _expand_concept_to_four_domain 函数 + Phase 0 fan-out 触发')
add_bullet('tests/qa-bench/questions.jsonl: D11-D15 共 5 道 + tools_must_all 检测器')
add_bullet('5/5 PASS（D11-D15 全 4 域覆盖）')
add_bullet('5 条新铁律（触发严格 == EXPLAIN_CONCEPT / 复用 _build_plan_step_input 0 修改 / 保留 LLM 已 planned 的非 4 域 tool / 必须独立 feature flag / prompt + 代码双重保险）')

add_h3('3.1.10 KB 页面"5 个统计全 0 + 暂无知识条目"修复（4 commit）')
add_bullet('根因 1：filterSourceType SPA 内存残留 = "auto_expansion"')
add_bullet('根因 2：健康度摘要的 entity/hypothesis/formula total 从未被主动 fetch')
add_bullet('根因 3：Service Worker 缓存了空 items 响应（CacheableResponsePlugin({statuses:[0,200]}) 只挡 5xx）')
add_bullet('方案 A：onMounted 重置 filter 状态 + chip 再点清除')
add_bullet('方案 B：SW BUMP VERSION + cacheWillUpdate 拒绝空 items')
add_bullet('方案 C：fetchKnowledge/fetchStats 加 try/catch + loadError')
add_bullet('方案 D：健康度摘要 onMounted 主动 fetch sub-entity total')

add_h3('3.1.11 声纹 sample_count 重置为 1')
add_bullet('DB 迁移 034_reset_voice_sample_count.py（down_revision=033_mvh）')
add_bullet('UPDATE members SET voice_sample_count = 1 WHERE voice_embedding IS NOT NULL（15 个已录入成员全部归零到 1）')
add_bullet('保留手动录入 +1 自增加权平均公式')
add_bullet('alembic 多 head 必须指定 target')

# --- 3.2 2026-06-29 ---
doc.add_page_break()
add_h2('3.2 2026-06-29 收官（11 项任务）')

add_h3('3.2.1 #043 账号持久化聊天历史 8 phase 完整收官（ChatGPT/Doubao 模式）')
add_bullet('Phase 1：ORM 模型 + alembic 039_chat_history.py（chat_sessions / chat_messages / chat_shares 三表）')
add_bullet('Phase 2：11 个后端 API 端点（17/17 e2e PASS）')
add_bullet('Phase 3：流式 chat 持久化修复（25/25 e2e PASS）')
add_bullet('Phase 4：前端 store 重构（chatHistory.ts + chatSessions.ts 同步）')
add_bullet('Phase 5：旧数据自动迁移（useChatMigration.js + localStorage chat_migrated_v1 标记 + 幂等键）')
add_bullet('Phase 6：UI 升级（SearchPalette / ShareDialog / ExportDialog / TagsEditor / SessionSidebar / MobileSessionDrawer / LongPressWrapper / MobileActionSheet / MobileSearchSheet）')
add_bullet('Phase 7：Celery 30 天清理（cleanup_soft_deleted_sessions_task + CHAT_HISTORY_RETENTION_DAYS=30）')
add_bullet('Phase 8：测试 + 沉淀（5 新测试文件：24+7+9+9+9 = 58 用例）')
add_bullet('12 条新铁律：流式持久化入场 append user / done 后落库 / CancelledError partial / JSONB flag_modified / best-effort / 迁移幂等 / 跨设备同步 / 软删除 30 天 / 越权防护 / localStorage 兜底 / 6 踩坑教训')
add_para('总耗时：22-30h，3 PR 收官', italic=True)

add_h3('3.2.2 v77 P2.6-G.2 模板批量管理页 /admin/templates（6 commits 收官）')
add_bullet('后端 schema 抽出 + list 端点 search/filter/pagination + batch toggle-active/delete endpoint')
add_bullet('前端桌面端 el-table + 批量操作（438 行 TemplatesView.vue）')
add_bullet('路由 + 侧边栏 + 移动 TabBar（query.tab=templates 跳转）')
add_bullet('Vitest 17 PASSED + Playwright B-17~B-20')
add_bullet('5 条新铁律：batch skip 知情提示 / builtin 允许 toggle 不删 / query 串分页 / dark 非 scoped 块 / el-table jsdom stub')

add_h3('3.2.3 v77 P2.6-F.5 builtin 一键复制 + is_active toggle UI（5 commits 收官）')
add_bullet('alembic 038 cloned_from_id 字段')
add_bullet('clone_template service + POST /meeting-templates/{id}/clone 端点')
add_bullet('builtin hover 复制按钮 + el-switch 启用/禁用')
add_bullet('6 条新铁律：builtin UX 闭环 / cloned_from_id 审计追溯 / disabled 双层防护 / 服务注释一致化')

add_h3('3.2.4 v77 P2.6-G.1 移动端 long-press 操作菜单入口（3 commits 收官）')
add_bullet('LongPressWrapper 包裹 + MobileActionSheet 弹菜单 + 桌面端 v-if 隐藏')
add_bullet('10 Vitest + 4 Playwright M-13~M-16')
add_bullet('5 条新铁律：long-press 是 hover 等价物 / LongPressWrapper 包裹整个卡片 / 移动端删除用 ElMessageBox.confirm / 桌面端 v-if 隐藏 / mobileActions callback 直接 emit')

add_h3('3.2.5 v77 P2.6-F.4 custom template-card hover 编辑/删除按钮（4 commits 收官）')
add_bullet('零 CSS 改动（meeting-view.css:304-318 已就绪）')
add_bullet('emit trick 复用 P2.6-F.3 save-template + el-popconfirm 二次确认')
add_bullet('4 Vitest + 2 Playwright B-11/B-12')
add_bullet('5 条新铁律：CSS 先行原则 / click.stop 嵌套 / el-popconfirm destructive / emit 多用途 / Playwright destructive test 用取消')

add_h3('3.2.6 v77 P2.6-F.3 MeetingTemplateDialog UI 入口闭环')
add_bullet('4 commits chain: MeetingCreateDialog 加按钮 + MeetingView 接事件 + 4 单测 + Playwright 集成 (B-05/B-06 从 skip → PASS)')
add_bullet('5 条新铁律：emit 定义必触发 / 双重防御 / editingTemplate trick / dialog nth(1) selector / 测试覆盖核心+Vitest 兜底')

add_h3('3.2.7 Ocean 主题按钮白字真 Bug 修复（Playwright 发现，commit 356a2bb0）')
add_bullet('Element Plus :root 比 variables.css 加载晚导致 --el-color-primary 被覆盖')
add_bullet('方案 C：[data-accent] 块内显式设 --el-color-primary + --el-color-primary-light-9')
add_bullet('4/4 Playwright PASS, 对比度 4.83:1 WCAG AA')

add_h3('3.2.8 5 件套视觉收官延伸（11 commits）')
add_bullet('KnowledgeToolbar 4 按钮：.btn-text utility class 同名冲突')
add_bullet('MemberView 录入声纹 ghost primary：variables.css 加 default + [data-accent] 双块规则')
add_bullet('VoiceprintView 波形颜色不一致：per-card max 归一化 + min alpha 0.12 + NaN 守卫')
add_bullet('SettingsView Hero 跟随主题：non-scoped [data-theme=dark].hero-bg source 顺序靠后赢 cascade')
add_bullet('VoiceprintEnrollFlow mobile icon + 5 处 transition token + webhint devDep')
add_bullet('nginx HSTS server-block + gzip_types 扩展：strict-transport-security 12→0 errors/route + 6 个 MIME')

add_h3('3.2.9 Knowledge 卡 analysis_status 卡 analyzing 真 bug 修复（commit 3653890b）')
add_bullet('Step 7 _reset_multimodal_data 无条件覆盖终态')
add_bullet('加 reset_status=False 参数 + Step 8 最终终态防御 + UI partial tag')

add_h3('3.2.10 webhint 二次扫描 + DB stuck 卡清理')
add_bullet('strict-transport-security 0 errors（9 路由全过）')
add_bullet('KB #14 #19（5 月预存 stuck 卡）验证 content 完整性 + UPDATE → done')

add_h3('3.2.11 声纹 strict merge 90% 识别率硬门禁（永久规则）')
add_bullet('scripts/verify_cross_meeting_recognition.py 跨会议验证')
add_bullet('scripts/restore_voiceprint.py 撤销单条 entry 但保留更早 merges')
add_bullet('< 90% → 立即 rollback（无条件，无需用户确认）')
add_bullet('90-95% → 报告用户决定')
add_bullet('≥ 95% → 接受')


# --- 3.3 2026-06-28 ---
doc.add_page_break()
add_h2('3.3 2026-06-28 收官（13 项任务）')

add_h3('3.3.1 v77 P2.6-F 收官 — transition: all token 化（1 commit）')
add_bullet('27 处 / 17 文件 → var(--transition-all-*)')
add_bullet('scripts/replace-transition-all-literals.js Node.js 脚本')
add_bullet('4 个 token：--transition-all-fast/normal/slow/slower')

add_h3('3.3.2 v77 P2.6-E 收官（3 commits）')
add_bullet('E.1 CSS-in-JS 收官：8 处 runtime :style → 7 个枚举 class + scss 55 → 105 行')
add_bullet('E.2 缓动字面量 token 化：70 处 → var(--ease-*) + 升级 --ease-out + 新增 --ease-quad')
add_bullet('E.3 KnowledgeView 拆分：1599 行 → 501 行（-68%），抽 5 个新组件到 components/knowledge/')
add_bullet('4 条新铁律：v-model 不能直接绑定子组件 props / el-pagination 改 :current-page + @current-change / Node.js word-boundary regex / 拆分巨型主 View 状态所有权')

add_h3('3.3.3 v77 P2.6-D 收官（4 commits）')
add_bullet('P2.6-D.1 PWA SW 强化：Background Sync (4 写场景) + Navigation Preload + Local Notification')
add_bullet('P2.6-D.2 动效治理收官：6 处重复 @keyframes 清理 + 3 --ease-* token + 12 --animation-* token')
add_bullet('P2.6-D.3 CSS-in-JS 收敛：3 处 avatar color → 枚举 class + scss 14 个枚举 class')
add_bullet('P2.6-D.4 Baseline 扩到 9 路由：desktop + mobile 各加 3 路由（projects/members/project-stats）')
add_bullet('4 条新铁律：PowerShell Set-Content -Encoding UTF8 写 UTF-8 BOM / Background Sync 仅适合幂等短写请求 / playwright baseline 必须 dev server 后台启 / token 化拆分渐进优于一次性铺开')

add_h3('3.3.4 v77 P2.6-C 收官 — EP 组件多主题透传补全 + Mobile Baseline 扩到 6 路由')
add_bullet('P0 三组件 75 条：el-tree / el-tree-select / el-date-picker / el-table 展开行 + 边框 + filter + sort')
add_bullet('P1 三组件 21 条：el-select 子级 / el-dropdown 子级 / el-tooltip / el-popover popper is-light/dark')
add_bullet('P2 五组件预留 47 条：el-cascader / el-transfer / el-autocomplete / el-color-picker / el-slider')
add_bullet('Mobile 登录态双注入修复：addCookies + addInitScript 写 localStorage（关键！router 守卫读这里）')
add_bullet('3 条新铁律：Playwright 登录态必须双注入 / EP 子选择器 dark 覆盖集中放 variables.css / EP dark 覆盖用 --el-*-bg-color 等 EP 内置 CSS 变量层')

add_h3('3.3.5 v77 P2.6-B 收官（4 子任务）')
add_bullet('子任务 1：Bug 修复（PaperHeader plain 按钮 dark 模式灰白）')
add_bullet('子任务 2.1：FallbackBlock dark 化（11/11 = 100% 收官）')
add_bullet('子任务 2.2：移动端 6 组件 dark 化（LongPressWrapper / MemberAvatar / PageHeader / MobileTaskCreateForm / SafeArea / MobileECharts）')
add_bullet('子任务 2.3：移动端 14 view dark 化（核心 3 + 中高 4 + 辅助 7）')
add_bullet('子任务 3：Playwright Desktop Baseline 6 路由（dashboard / chat / knowledge / tasks / meetings / settings）')
add_bullet('4 条新铁律：dark 模式跨组件覆盖必须非 scoped <style> 块（第 5 次强化）/ JS 端 ECharts 调色板用 getComputedStyle 读 token / Stylelint 禁用 hex 颜色 / Playwright baseline 接受 CI auto-commit 平台后缀差异')

add_h3('3.3.6 v77 P2.6-A 收官（commit 36049629）')
add_bullet('paper 14 组件 + 桌面 5 view + ChartBlock token dark 全面化')
add_bullet('移动端 9/15 → 15/15 + Rich Block 11/11 dark 化收官')
add_bullet('ChartBlock JS 端 getComputedStyle + MutationObserver')

add_h3('3.3.7 pgvector embedding truth value bug 修复（会议 64 报 500）')
add_bullet('not numpy_array 抛 "truth value ambiguous"')
add_bullet('改用 is None')
add_bullet('2 处生产代码修复 + 3 case 验证')

add_h3('3.3.8 SQLAlchemy JSONB flag_modified 教训（会议 64 polished mirror）')
add_bullet('JSONB 列内部元素 mutate commit() 静默不持久化')
add_bullet('必须 flag_modified(m, "field") 强制 UPDATE')

add_h3('3.3.9 AudioPlayer Infinity:NaN 修复（WebM 流式音频）')
add_bullet('audio.duration 初始 Infinity → 显示 "Infinity:NaN"')
add_bullet('加 duration prop + formatTime 防御 Number.isFinite + 后端预知时长')

add_h3('3.3.10 v76 视觉回归（Playwright Baseline）废弃决策')
add_bullet('最近 50 个 GH Actions run 20 个失败（40%）')
add_bullet('mock token 数据漂移 1-2% + OS suffix 跨平台陷阱 + desktop baseline auto-commit workflow bug')
add_bullet('完全废弃 CI 视觉回归 job（保留 spec 作本地 dev 用）')

add_h3('3.3.11 083 事件 + KMeans+声纹 Merge Pipeline 全栈优化')
add_bullet('P0 防护（sil_floor + cluster_centers 合并 + strict 策略）')
add_bullet('P1 配套（Hungarian 投票覆盖 + inject/inspect）')
add_bullet('P2（learn opt-in + cluster_id_history）')
add_bullet('083 杜同贺 86.7% → 100%')

add_h3('3.3.12 会议 #151 声纹循环净化')
add_bullet('strict 2/3 提升（王天志 583 / 张宏魁 125）')
add_bullet('1/3 BLOCKED（杜同贺 0.60 当前 embedding 不是本人）')
add_bullet('#135 override 错标诊断 + 8 条新铁律')

add_h3('3.3.13 qa-bench 5 轮迭代 39% → 84%')
add_bullet('100 题基线 + 75 拓展 + 240 拓展 + 495 动态 = 910 题')
add_bullet('知识库增长：64 → 247（+183 条, +286%）')

add_h3('3.3.14 Agent 质量 5 大根因修复（14 commits）')
add_bullet('根因 1：TOOL_REGISTRY 启动时未注册（commit d36d1db）')
add_bullet('根因 2：LLM 代理层 fake tool_call（commit d36d1db + e2a9a49）')
add_bullet('根因 3：get_member_profile dead import + is_active 过滤 alumni')
add_bullet('根因 4：长期记忆干扰（commit e2a9a49）')
add_bullet('根因 5：synthesis 阶段 fake XML 泄露（commit e2a9a49）')

# --- 3.4 2026-06-27 ---
doc.add_page_break()
add_h2('3.4 2026-06-27 收官（12 项任务）')

add_h3('3.4.1 智能论文阅读器 v26 + v26.1 回归修复（7 处回归 + 8 条新铁律）')
add_bullet('chemFormat Unicode 化（HTML 二次转义修复）')
add_bullet('正则 {0,N}? + |$ 是经典陷阱，禁止使用（铁律 2）')
add_bullet('OCR 半截 JSON 必须有专门 _stripMultimodalBlocks 函数处理（铁律 3）')
add_bullet('_escapeHtml 二次转义是隐形杀手（铁律 4）')
add_bullet('图片必须按"真实图号 + isCoreFigure 分级"插入（铁律 5）')
add_bullet('阅读器宽度必须有"主列 + 工具栏"分层（铁律 6）')
add_bullet('内联图片必须有"高置信度"三重校验（铁律 7）')

add_h3('3.4.2 v70-v76 前端字面色 token 化 + 视觉回归测试收官')
add_bullet('~340 处 hex 字面色 → var(--color-*) token，dark mode 全面修复')
add_bullet('stylelint 字面色禁用 + 组件级 CSS 测试')
add_bullet('Playwright 视觉回归 5 件套')

add_h3('3.4.3 pre-commit hook auto-add web/dist/（CLAUDE.md 教训第 4 次沉淀后兜底）')
add_bullet('scripts/check-dist-before-commit.sh 业务逻辑')
add_bullet('.git/hooks/pre-commit 薄 wrapper')
add_bullet('检测 staged web/src/ 改动 + 本地新 dist hash 产物')
add_bullet('5 条铁律：改 web/src/ 必跑 npm run build + commit 前必看 dist / hook 触发条件 / hook 不是 hard block / 测试 commit 触发的 post-commit 自动 push 必须警惕 / 新成员必须手动 setup hook')

add_h3('3.4.4 声纹 sample_count 重置为 1')
add_bullet('alembic 034_reset_voice_sample_count.py')
add_bullet('保留手动录入 +1 自增 + 自动学习链已删除')

add_h3('3.4.5 声纹 batch bug 修复（推到主路径）')
add_bullet('ERes2Net_aug.py:__extract_feature 强制 batch=1 导致 97% 沉默失败')
add_bullet('ThreadPoolExecutor(8) + Lock 修复后 100% 段有效')

add_h3('3.4.6 会议发言人重处理流程标准化（reprocess_meeting.py）')
add_bullet('9 步 CLI（load / extract / cluster / vote / assign / backup / apply / regen / verify）')
add_bullet('文件备份（不是 DB schema 备份）')
add_bullet('主机端 wrapper (.ps1/.bat)')
add_bullet('11 条铁律')

add_h3('3.4.7 深度学习提升方向 125+ 项')
add_bullet('A/B/C 三类 + 智能体 + 备用资产 + GPU 调度')
add_bullet('6 个月路线图，PyTorch > TF，跳过 MATLAB')

add_h3('3.4.8 v77 P2.6-B 收官')
add_bullet('Bug 修复 + 移动端 14 view + 6 组件 + 1 Block dark 化 + Desktop Baseline 6 路由')

add_h3('3.4.9 v77 P2.6-A 收官')
add_bullet('paper 14 组件 + 桌面 5 view + ChartBlock token dark 全面化')

add_h3('3.4.10 屏幕浮窗软件识别（FPS 是关键鉴别信号）')
add_bullet('右上角 GPU/CPU/FPS/延迟 浮窗是 NVIDIA App Performance Overlay')
add_bullet('HWiNFO64 自动化局限')

add_h3('3.4.11 m4a 录音处理全链路优化')
add_bullet('2.9h m4a 处理 + celery-worker GPU 化 + batch voiceprint 24 倍加速')
add_bullet('4 条 KB 入库')

add_h3('3.4.12 Docker Desktop 修复与数据迁移')
add_bullet('WSL docker-desktop-data 重建 + junction E 盘化')
add_bullet('aliyun 正确路径 + 清华 pip 重试')
add_bullet('.dockerignore 17 倍提速（12GB → 700MB）')
add_bullet('frp 用户级 AtLogOn')

# --- 3.5 2026-06-26 ---
doc.add_page_break()
add_h2('3.5 2026-06-26 收官（11 项任务）')

add_h3('3.5.1 v69 桌面端 dark mode 全面重构（3 阶段：P0+P1 已完成，P2 待做）')
add_bullet('P0 视觉修复（commit 71bb394a）：5 tokens + 14 EP 覆盖 + MainLayout + Dashboard')
add_bullet('P1a 多主题切换基建（commit 55865fe2）：6 套主题（3 主色 × 2 明暗）')
add_bullet('P1b 10 桌面视图 dark 适配（commit 7e0976d8）：ChatViewSSE / TaskView / TaskTrash / MeetingView / MeetingDetailView / KnowledgeView / KnowledgeDetailView / ProjectView / MemberView / AgentTracesView')

add_h3('3.5.2 v31.2.6 login_limiter Redis 化 + Retry-After 响应头')
add_bullet('login_limiter = AsyncRedisRateLimiter(max_attempts=5, window_seconds=300)')
add_bullet('FastAPI HTTPException(headers={"Retry-After": str(window_seconds)})')
add_bullet('middleware 路径额外补 dict(e.headers) 转发')
add_bullet('login: 前缀与 middleware tier 命名一致')

add_h3('3.5.3 v31.2.5 启用 AsyncRedisRateLimiter 替换 RateLimiter（抗 docker restart）')
add_bullet('5 个 tier 实例全部从 RateLimiter 改为 AsyncRedisRateLimiter')
add_bullet('middleware 3 处同步调用 await 化')
add_bullet('silent degradation try/except 兜底 Redis 故障')
add_bullet('4 条铁律：check + record 必须分开 / uvicorn 写响应头是小写 / socket 流式响应读响应头必须用 \\r\\n\\r\\n / Redis 限流"抗重启"必须 ZSET 持久化')

add_h3('3.5.4 v31.2.4 AsyncRedisRateLimiter 类实现 + memory 沉淀 + by_user dashboard')
add_h3('3.5.5 v31.2.3 rate-limit 基建收尾（X-RateLimit-Policy 头 + SSE tier + auth prefix 匹配）')
add_bullet('3 条铁律：限流响应头必须有 tier 信息 / SSE 长连接必须独立 tier / 路径前缀匹配用 startswith 而非 in path')

add_h3('3.5.6 v31.2.2 rate-limit 进阶强化（regex 精确路径 + user_id 维度限流）')
add_bullet('regex 锚定 ^...$ + 路径分隔')
add_bullet('_try_attach_user_id middleware 顶部注入 user_id 到 request.state')
add_bullet('2 条铁律：substring 路径匹配必须用 regex 或 prefix / 限流 key 注释必须真实')

add_h3('3.5.7 v31.2.1 rate-limit 边界强化（XFF 空 IP 兜底 + auth/analytics 嵌套防御）')
add_bullet('XFF 首段为空（", 1.2.3.4" / "   " / ",,,,,"）时，统一兜底为 "unknown"')
add_bullet('方案 B1：/analytics 分支前置守卫 not path.startswith("/api/v1/auth/")')

add_h3('3.5.8 v31.3 Whisper 常驻 + 推理加速')
add_bullet('模型常驻 GPU 8GB（取消 lazy load + 10 分钟空闲卸载）')
add_bullet('flash_attention=True（Blackwell sm_120 当前不支持，ctranslate2 4.8.0 报错）')
add_bullet('实际部署 flash_attention=False，等上游补 sm_120 flash attn 2 内核')
add_bullet('3 条铁律：18s 冷启动 vs 8GB 常驻——用户决策优先级 / flash_attention=True 不加速加载，只加速推理 / files= 文档有但不能用')

add_h3('3.5.9 v31.3.1 whisper 容器 bind mount（解决"Dockerfile COPY 烧镜像"陷阱）')
add_bullet('Dockerfile.whisper 删 COPY app/whisper_server.py')
add_bullet('docker-compose.yml 加 - ./app/whisper_server.py:/app/whisper_server.py:ro')
add_bullet('3 条新铁律：Dockerfile COPY 源码是反模式 / debug print 必须放在执行路径上 / 容器 bash -c "head file" 时 docker exec 有 cwd 解析 bug')

add_h3('3.5.10 webhook deploy 链断裂修复')
add_bullet('5 真相：webhook service 收到 200 OK 但 deploy 失败')
add_bullet('git@github.com: Permission denied (publickey) — server-side deploy key 与 GitHub 端不匹配')
add_bullet('5 铁律：webhook 失败根因必须先看 /var/log/webhook-deploy.log / 公钥定期轮换 / .env.webhook 守护 / 本地 SSH 不到 server 时只能从公网探测 / 60s 不更新 = 异常')

add_h3('3.5.11 8 个 webhint PWA 警告全栈修复（commit 08f440f + c855f0e）')
add_bullet('Nginx 缺 .webmanifest MIME → server block 加 types { application/manifest+json webmanifest; }')
add_bullet('vite-plugin-pwa 输出 manifest 不带 hash → Vite plugin manifestHashPlugin closeBundle 钩子')
add_bullet('/registerSW.js 静态注入无法 cache-busting → injectRegister: null + useRegisterSW')
add_bullet('删除 manifest.webmanifest 后 SPA fallback 误返 index.html → location = /manifest.webmanifest { return 410; }')

# --- 3.6 2026-06-25 ---
doc.add_page_break()
add_h2('3.6 2026-06-25 收官（4 项任务）')

add_h3('3.6.1 v31.2.x rate-limit 4 版本收尾')
add_bullet('v31.2.1/2/3/4 共 7 条铁律 + 4 个 verify 脚本（verify_v31_2_1_xff_empty / verify_v31_2_1_nested_path / verify_v31_2_2 / verify_v31_2_3）')

add_h3('3.6.2 v31.2 rate-limit IP 维度 + Redis 持久化')
add_bullet('滑动窗口算法 + Redis ZSET')

add_h3('3.6.3 会议 #95 声纹重置 + 重识别教训')
add_bullet('6 条铁律：psycopg2 transaction rollback 静默吃 UPDATE / speaker_mapping 完全错标时必须用 KMeans 重聚类 / Whisper 幻觉段不能用作声纹学习 / speaker_mapping 与 meeting_participants 必须互相对齐 / 用户认知 vs 数据库不一致时必须先汇报数据 / 重置声纹是一次性单向操作')

add_h3('3.6.4 edge-tts 6.1.9 TrustedClientToken 过期 → TTS 500（commit 41cf204）')
add_bullet('edge-tts 6.1.9 已失效，Microsoft 返回 403 Forbidden')
add_bullet('升级到 edge-tts 7.2.8（PyPI 最新版）')
add_bullet('requirements.txt 不能盲目锁 == 版本，用 >=X,<Y 范围')

# --- 3.7 2026-06-24 ---
add_h2('3.7 2026-06-24 收官（1 项任务）')

add_h3('3.7.1 sentence-transformers 5.6.0 升级（Phase 1+2 收官）')
add_bullet('CLAUDE.md 标"❌ sentence-transformers 升级（未做）"')
add_bullet('实测后 100% 完成 + 超额（原 plan 担心跨 3 大版本破坏性，0 破坏，qa-bench 38% → 42% 反升）')
add_bullet('跳过 Phase 3：ONNX 实测 GPU 上慢 12-22x（反优化），保持 torch/GPU')
add_bullet('5 大铁律：清华源限速 torch 2.12+ / docker compose build 别用环境变量 HTTPS_PROXY 污染全局 / ONNX backend 在 GPU 上是反优化 / sentence-transformers 升级时分 Phase / ST 5.6.0 的 Pooling 现在支持 include_prompt')

# --- 3.8 2026-06-22 ---
add_h2('3.8 2026-06-22 收官（1 项任务）')

add_h3('3.8.1 深度学习提升方向 125+ 项')
add_bullet('A/B/C 三类 + 智能体 + 备用资产 + GPU 调度')
add_bullet('6 个月路线图')
add_bullet('PyTorch > TF，跳过 MATLAB')

# --- 3.9 2026-06-20 ---
doc.add_page_break()
add_h2('3.9 2026-06-20 收官（v28 论文图片结构化字段后端集成）')

add_h3('3.9.1 v28 step 2 + step 3 — 8 条铁律')
add_bullet('Vision LLM 看不到图外 caption，figure_no 覆盖率仅 20%')
add_bullet('单图并发跑 2 个 LLM 调用 = 共用 semaphore 池')
add_bullet('return_exceptions=True 必须配独立异常处理')
add_bullet('anchor_paragraph_index 用 \\n\\n + [PAGE:N] 简单分段落')
add_bullet('anchor_text fallback regex 必须容忍 "Fig" / "Figure" / "Figs" / "Fig." 4 种写法')
add_bullet('SQLAlchemy Column 重名 Index 在 ORM 里冲突')
add_bullet('v28 字段写入必须容错每一种字段类型')
add_bullet('volume mount 新 alembic 文件必须 docker cp + 清 __pycache__')

add_h3('3.9.2 v28 step 4 — paperAdapter 简化为读后端字段')
add_bullet('后端 schema + API 补全 12 字段')
add_bullet('前端 _normalizeImages 透传 12 字段 + _buildFigureRegistry 简化')
add_bullet('Graceful Degradation：vision 字段全 null 时回退旧逻辑')
add_bullet('7 条铁律：删前端推断逻辑前先看后端 schema / snake_case → camelCase 字段映射必须一对一完整 / anchor_text fallback 弥补 vision 模型 20% 覆盖率 / === true 严格比较布尔字段 / datetime 字段序列化必须 str() 转字符串 / 测试 fixture 必须反映真实数据形状 / Legacy 函数重命名要带后缀')

add_h3('3.9.3 v28 step 5 — RightImageRail sectionHint 精准推荐')
add_bullet('KnowledgeDetailView 独立接 IO + currentSectionFigures 重写')
add_bullet('9 条铁律：v27.2 注释不代表已实现 / RightAnchorNav 的 IO 不应被 KnowledgeDetailView 依赖 / sectionHint 匹配算法必须按核心词重叠 / setupSectionObserver 必须在 paper.value 赋值之后 + await nextTick() / IntersectionObserver 必须 disconnect + null 避免内存泄漏 / activeSectionId.value 变化判断必须加 !== bestId / rootMargin 与 RightAnchorNav 保持一致 / fallback 永远保留 / onMounted 在 KnowledgeDetailView 不能重复')

add_h3('3.9.4 v28 step 6 — 内嵌图 confidence ≥ 0.85 阈值')
add_bullet('PaperSectionRenderer 加 showHighConfidenceOnly prop（默认 true） + HIGH_CONFIDENCE_THRESHOLD = 0.85')
add_bullet('ReadingToolbar 加"高质量图"切换按钮 + maxConfidencePct computed + emit toggle-high-confidence')
add_bullet('6 条铁律：confidence 阈值常量必须定义在 props 同模块 / 低置信度图默认隐藏 / 徽章只在 showInlineFigures=true 时显示 / localStorage 持久化默认值必须明确处理 / 徽章 v-if=showHighConfidenceOnly 而非 always 显示 / filter 顺序 showInlineFigures → showHighConfidenceOnly')

add_h3('3.9.5 v28 step 7 — 4 篇真实 PDF 端到端验证')
add_bullet('37 张图 100% vision 覆盖 + 100% publisher 准确 + mean conf 0.93')
add_bullet('figure_no 覆盖率 11%（vision 模型 20% 上限）')
add_bullet('2 bug 修复：_compute_anchor_for_images 不识别中文边界 / NoneType 错误')
add_bullet('5 铁律：Python \\b 不处理中文边界 / vision model 输出不稳定 / 验证脚本必须能识别 vision model 输出异常 / 测试脚本必须能跑在容器内（DB 用 db host） / _reset_multimodal_data 会清空所有 vision 字段')

add_h3('3.9.6 v28 step 8 — IO Hysteresis + rAF 节流')
add_bullet('ACTIVATE_THRESHOLD = 0.35 + HYSTERESIS_LOWEST = 0.15')
add_bullet('rAF 节流：60 fps → 16ms 一次')
add_bullet('visibilityMap 跨 route 清空')
add_bullet('8 滚动场景验证全过')

# --- 3.10 2026-06-19 ---
doc.add_page_break()
add_h2('3.10 2026-06-19 收官（Phase 7 多模态知识库 + 声纹 batch bug 修复）')

add_h3('3.10.1 Phase 7 多模态知识库（图片/公式/表格 OCR 入库）')
add_bullet('2 张新表：knowledge_images + knowledge_extractions（统一 formula/table/chart/image_block）')
add_bullet('OCR 服务抽象层（app/services/ocr_service.py）：主后端 LLM-Vision 复用 vision_service，可选 Tesseract 备选')
add_bullet('多模态解析管线（app/services/multimodal_extraction_service.py）：PDF/PPTX 提取嵌入图片 → 缩放 → MinIO → asyncio.Semaphore(4) 并发 OCR')
add_bullet('3 个新 API：GET /knowledge/{id}/images + GET /knowledge/{id}/extractions + POST /knowledge/{id}/extract-multimodal')
add_bullet('5 个新 settings：MULTIMODAL_OCR_BACKEND/_CONCURRENCY=4/_MAX_IMAGES_PER_DOC=20/_MAX_IMAGE_PIXELS=2.5MP/_MIN_IMAGE_PIXELS=10k')
add_bullet('2 个新前端组件：KnowledgeImageGallery.vue + KnowledgeExtractionsPanel.vue')
add_bullet('8 条铁律：多模态入库拆 2 张表 / OCR 后端选 LLM-Vision 复用 vision_service / 并发控制 asyncio.Semaphore / 图片处理前置过滤 + 缩放 / session 隔离 / 列表接口不能 mutate ORM 对象 / vision 模型会泄露 ThinkingBlock 序列化字符串 / alembic 链断裂必须 1 行修复')

add_h3('3.10.2 会议发言人重处理流程标准化（reprocess_meeting.py）')
add_bullet('9 步 CLI（load / extract / cluster / vote / assign / backup / apply / regen / verify）')
add_bullet('11 条铁律：ERes2Net 强制 batch=1 / SQLAlchemy 静默忽略未映射属性 / 备份必须独立于 DB schema / regen 必须可独立运行 / CLI step 自动依赖解析 / 短段 (< 0.6s) 标"发言人?"是合法状态 / 527 段保留 "发言人?" 不影响前端体验 / 8 字段 verify 一行 SQL 必查 / 应到会人数影响 KMeans K 搜索范围 / LLM 重生成必须用 ANALYSIS_PROMPT / MeetingParticipant 表必须同步更新')

add_h3('3.10.3 声纹 batch bug 修复（推到主路径）')
add_bullet('ERes2Net_aug.py:__extract_feature 强制 unsqueeze(0) 折叠 batch')
add_bullet('ThreadPoolExecutor(8) 并行单条 + threading.Lock 保护 pipeline.model')
add_bullet('7 条铁律：上游库的 bug 必须 app 层绕开 / 所有会议识别质量改进要 push 到主路径 / ThreadPoolExecutor + Lock 组合 / 声纹 embedding 验证不能只看长度 / 沉默失败比明显错误更可怕 / 重启后端是 volume 挂载的硬要求 / 100% 段有效是默认值')

# --- 3.11 2026-06-18 ---
doc.add_page_break()
add_h2('3.11 2026-06-18 收官（移动端 10 PR + 三连环修复）')

add_h3('3.11.1 移动端 10 PR 全栈定制收官（commit 9026c07）')
add_bullet('18 个移动端页面 + 12 个移动端组件 + 4 个 PWA 离线策略')
add_bullet('路由级双栈架构：useIsMobile.js 判定 + resolveMobile.js 动态 import')
add_bullet('PWA 4 策略：manifest + service worker（workbox）预缓存 app shell + useSafeArea 读 iPhone 安全区 + 离线 IndexedDB 兜底')
add_bullet('视觉回归测试：Playwright 5 viewport × 13 核心页面，CI 截图对比基线')

add_h3('3.11.2 EP useOrderedChildren.removeChild null 崩溃修复（commit f8d27015）')
add_bullet('Vite plugin transform 阶段 patch EP 源码')
add_bullet('5 条铁律："上游已知 bug 但未修复"系列再次扩展 / EP source use-ordered-children/index.mjs 是 pane 注册中心的唯一源 / patch 文件 ID 用完整 node_modules 路径正则 / 升级 EP 大版本时 plugin 自动失效但要 warn loud / dev mode 不 patch')

add_h3('3.11.3 桌面"正在听会"指示器不接进度修复（commit f099e7e5）')
add_bullet('新建 MeetingRoomView 全屏页镜像移动端 MobileMeetingRoom')
add_bullet('router fallback 改正 + MeetingView.resumeRecording 改 navigate')
add_bullet('4 条铁律：同一功能必须桌面/移动 UX 一致 / resolveMobileComponent 必须给桌面 fallback 真正的桌面组件 / 改动多处复用页面结构时 grep 桌面 + 移动两个版本 / 新建视图文件优先放 web/src/views/ 根')

add_h3('3.11.4 三连环修复（EP patch + MeetingRoomView + /auth/me 限流）')
add_bullet('事故链：本地只 commit 没 push，误以为是 webhook 链断')
add_bullet('新铁律：commit 后必 verify push 真到 origin / 怀疑 webhook 断时第一步先看 origin/main / /auth/ 下路径必须分级限流 / docker compose v1 vs v2 服务器上不互通 / CLAUDE.md 752 行铁律对后端依然适用')

# --- 3.12 2026-06-17 ---
doc.add_page_break()
add_h2('3.12 2026-06-17 收官（部署与基础设施重建）')

add_h3('3.12.1 Docker Desktop 引擎崩溃 + 镜像源治理 + 数据 E 盘化')
add_bullet('WSL2 docker-desktop-data 发行版丢失 → com.docker.service 7-9 分钟反复启停')
add_bullet('C 盘 24GB Docker 缓存清空 + 数据全 E 盘化（junction 透明重定向）')
add_bullet('huaweicloud 镜像源 404 → aliyun 正确路径（Debian bookworm-security 走 debian-security/ 独立路径）')
add_bullet('aliyun PyPI 限速 600KB/s → 清华源 + pip --retries 10 --timeout 60')
add_bullet('新增 .dockerignore（build context 12GB → 700MB 17 倍提速）')
add_bullet('10 条铁律：junction 透明重定向 / WSL Docker 引擎恢复流程 / Dockerfile 镜像源选择 / PyPI 限速真相 + pip 重试 / apt-get install 必加 fallback / .dockerignore 是必须的 / docker-compose.override.yml 默认加载 / frp 客户端 Windows 开机自启 / dockerproxy.net 500 错误 / ~/.docker/config.json 不要随便加 proxies 字段')

add_h3('3.12.2 webhook deploy 链断裂修复（commit c9c60ca6）')
add_bullet('5 真相：webhook service 收到 200 OK 但 deploy 失败')
add_bullet('git@github.com: Permission denied (publickey)')
add_bullet('deploy-auto.sh 加 .env.webhook 守卫 + 重新生成 ed25519')

# --- 3.13 2026-06-15 ---
doc.add_page_break()
add_h2('3.13 2026-06-15 收官（任务提醒 v2 + Agent 质量 + LLM 元话语修复）')

add_h3('3.13.1 任务提醒体系 v2 全面优化（commits 223ea74 + ba75e32）')
add_bullet('所有 reminder 统一在 11:00 AM 北京时间窗口发送（±60min 容差）')
add_bullet('任何微信消息 = ack 取消该用户所有 pending 提醒')
add_bullet('新增 acknowledged 状态 + acknowledge_all_user_reminders 服务方法')
add_bullet('v2.1 二次简化：用户原话"用户发任何内容都是不再提醒"')
add_bullet('状态机：pending → sent / acknowledged / cancelled')
add_bullet('v2 漏修补救（commit d0ddf49e）：proactive scheduler 也必须走 11AM 窗口')

add_h3('3.13.2 Agent 质量 + qa-bench 闭环（5 大根因修复）')
add_bullet('5 大根因：TOOL_REGISTRY 启动时未注册 / LLM 代理层 fake tool_call / get_member_profile dead import + is_active 过滤 alumni / 长期记忆干扰 / synthesis 阶段 fake XML 泄露')
add_bullet('5 格式解析（agentic_loop._parse_fake_tool_calls）：Mistral/Qwen / Anthropic legacy / 简化 / 裸 JSON / 混合格式')
add_bullet('schema-aware alias（_normalize_fake_tool_input）：name → member_name 按 Pydantic model_fields 反查自动映射')
add_bullet('关键改进：search_knowledge 返回 0 结果时 hint / 数据缺失警告 / intent_classifier 增强')
add_bullet('前端 UI 干净化：useUiStore.js + ChatViewSSE.vue + RichContent.vue')

add_h3('3.13.3 LLM 元话语/thinking 文本泄露修复（双管齐下）')
add_bullet('prompts.py 硬规则（19 种元话语前缀 + 正反例）')
add_bullet('后端 _strip_meta_thinking 兜底剥除（19 个 regex pattern + while 循环）')
add_bullet('SSE done 事件带 text_without_json')
add_bullet('前端 useChatStream.ts done 时替换')

add_h3('3.13.4 Rich Block 统一包装铁律（杨慈是谁呀"暂无成员"修复）')
add_bullet('Rich Block 提取两条路径必须保证 data 形态统一：fake tool_call 路径 / JSON 段路径')
add_bullet('修复模式：chat_engine.py:373-385 对 rb_type=="member" 统一包装为 {members: [...]}')

add_h3('3.13.5 移动端"声纹测试"真识别测试修复（commits de7ef8aa + 22d5570a）')
add_bullet('移动端 VoiceTestFlow.vue 状态机升级为 5 态（idle → recording → recorded → testing → result）')
add_bullet('7 铁律："X 测试"组件先 grep 桌面端有没有同名实现 / 状态机切换不能在 stopTest 里 / iOS Safari MediaRecorder 兜底 / 不要手动设 Content-Type / 5 秒自动停止 + 5 步步骤展示 / 错误降级构造 steps / a11y 4 属性套件')

add_h3('3.13.6 v-model 命名必须跟组件 prop 名严格匹配（commit f84524cf）')
add_bullet('v-model 必须对应组件 prop 名 modelValue（默认）')
add_bullet('v-model:foo 必须对应组件 prop 名 foo')
add_bullet('Vue 不会编译报错，调试"点击没反应"类问题第一步 grep v-model:xxx')

# --- 3.14 2026-06-14 及更早 ---
doc.add_page_break()
add_h2('3.14 2026-06-14 及更早收官')

add_h3('3.14.1 2026-06-14 GitHub Actions 3 jobs 跑通期间踩的 6 个真实坑')
add_bullet('坑 1：npm install 异常中断 → lock 与 package.json 不同步 → CI EUSAGE')
add_bullet('坑 2：workflow paths filter 必须包含 lockfile')
add_bullet('坑 3：改 workflow yml 自身不触发该 workflow run')
add_bullet('坑 4：workflow_dispatch / push / pull_request 是 3 个独立触发器')
add_bullet('坑 5：GitHub Actions bot 默认 read-only，push 需要 permissions: contents: write')
add_bullet('坑 6：Playwright baseline 按 OS suffix 区分（*-linux.png / *-win32.png / *-darwin.png）')

add_h3('3.14.2 2026-06-12 v4 收官（34 工具 @tool 装饰器 + 12 类 Rich Block + 真实 SSE 流式 + 10 字段响应）')

add_h3('3.14.3 2026-06-12 方案 C：Agent 单阶段流式渐进综合架构（6 stage 收官）')
add_bullet('取消 brief/detail 双层 → 单阶段流式综合（intent → agentic_loop → critique → done）')

add_h3('3.14.4 2026-06-12 chat_engine_legacy 30 天承诺提前 15 天收官（2026-06-29, commit 817f1ffa）')

add_h3('3.14.5 Phase 1-6 全部完成 + 160+ 测试全过')

add_h3('3.14.6 2026-06-11 webhint 优化纪律 + 11 个 webhint 警告全栈修复（commit 08f440f + c855f0e）')

add_h3('3.14.7 2026-06-10 git add -A 静默跳过 .gitignore 内文件教训（CLAUDE.md 第 4 次沉淀）')

add_h3('3.14.8 2026-06-09 nginx 扫描器正则误杀 /webhook 修复 + frp 0.x 重连逻辑')

add_h3('3.14.9 2026-06-08 Webhint 优化纪律（11 警告全栈修复）')

add_h3('3.14.10 2026-06-06 会议纪要标准格式（2026.5.28 例行例会 信息密度）')

add_h3('3.14.11 2026-06-04 代码质量规范（API 层 / 前端架构 / 测试规范 全面升级）')

add_h3('3.14.12 2026-06-03 faster-whisper 7 层过滤（后端 + whisper_server.py + HALLUCINATION_STRONG/WEAK）')

add_h3('3.14.13 2026-06-02 声纹会议系统全面修复（8 commits）')


# ============================================================
# 第四部分：沉淀的核心铁律
# ============================================================
doc.add_page_break()
add_h1('四、沉淀的核心铁律（高频复用）')

add_h2('4.1 部署铁律')
add_para('CLAUDE.md 752 行铁律（项目最高频引用）：', bold=True)
add_bullet('任何后端 Python 文件改动必跑 docker compose restart app celery-worker（volume 挂载只换文件不换模块缓存）', bold=True)
add_bullet('任何改 web/src/ 必跑 npm run build + git add -f web/dist/（pre-commit hook 兜底）', bold=True)
add_bullet('修改 nginx 后必跑 6 点 curl 验证（HTML + CSS + JS + PNG + manifest + sw.js）', bold=True)
add_bullet('任何 alembic 迁移必 docker cp + docker exec alembic upgrade + restart app', bold=True)

add_h2('4.2 数据铁律')
add_bullet('KB 数据清理必须按 source_type 隔离自动/真实用户（防御性 WHERE）', bold=True)
add_bullet('删除数据任务必须三段式（scan → 人审 → apply + --confirm）+ JSON 文件备份', bold=True)
add_bullet('声纹 strict merge 90% 识别率硬门禁：< 90% 自动 rollback', bold=True)
add_bullet('SQLAlchemy session.rollback() 不撤销 committed row，必须显式 fixture teardown delete', bold=True)

add_h2('4.3 CSS / 前端铁律')
add_bullet('dark 模式 + 跨组件覆盖必须非 scoped <style> 块（v60-v67 教训第 8 次强化）', bold=True)
add_bullet('CSS 变量必须放在 variables.css 末尾，新代码不写 hex 字面色', bold=True)
add_bullet('--transition-all-* token + --ease-* token + --animation-* token 收敛字面量', bold=True)
add_bullet('CSS-in-JS runtime :style → 枚举 :class 收敛 8 类', bold=True)
add_bullet('EP 子选择器 dark 覆盖集中放 variables.css（不分散到各组件）', bold=True)
add_bullet('el-table 在 jsdom 测试中必须用 global.stubs 阻止 slot 渲染', bold=True)

add_h2('4.4 后端铁律')
add_bullet('typing import CI 检查：新文件 Dict/List/Optional 必须有 from typing import ...', bold=True)
add_bullet('SSE 事件 delta 语义显式标注 [increment] 或 [snapshot]', bold=True)
add_bullet('JSONB mutate 必须 flag_modified（CLAUDE.md 2026-06-28 教训）', bold=True)
add_bullet('fastapi 0.125+ HTTPBearer auto_error=True 默认 401（不是 403）', bold=True)
add_bullet('ToolContext 注入（跨 event loop 安全，方案 C 铁律 1）', bold=True)
add_bullet('background sync 仅适合幂等短写请求（SSE/WS/大文件不适合）', bold=True)

add_h2('4.5 限流铁律（v31.2.x 收官）')
add_bullet('substring 路径匹配必须用 regex 或 prefix（v31.2.1/2 教训）', bold=True)
add_bullet('限流 key 必须有 fallback 值（空字符串 key 表面看"无 IP"实际是有效共享 key）', bold=True)
add_bullet('SSE 长连接必须独立 tier（v31.2.3 教训）', bold=True)
add_bullet('限流响应头必须有 tier 信息（X-RateLimit-Policy）', bold=True)
add_bullet('Redis 限流"抗重启"必须 ZSET 持久化（v31.2.5 教训）', bold=True)
add_bullet('check + record 必须分开（v31.2.5 教训，composite 函数会污染计数）', bold=True)
add_bullet('FastAPI HTTPException(headers={"Retry-After": str(window_seconds)}) 标准机制（v31.2.6 教训）', bold=True)

add_h2('4.6 声纹 / 会议铁律')
add_bullet('ERes2Net_aug.py:__extract_feature 强制 batch=1 → 必须 ThreadPoolExecutor + Lock 修', bold=True)
add_bullet('Whisper 反幻觉必须三层防护：condition_on_previous_text=False + no_speech_threshold=0.6 + HALLUCINATION_STRONG/WEAK', bold=True)
add_bullet('HARDCODED_ALIASES 与 PHONETIC_CORRECTIONS 必须单源', bold=True)
add_bullet('清洗 hook 必须早于 LLM polish（防 LLM 把错人名写进 key_points/decisions）', bold=True)
add_bullet('Fuzzy 阈值 ≤ 1 编辑距离（不能放宽到 2）', bold=True)
add_bullet('speaker_mapping 与 meeting_participants 必须互相对齐（重识别时同步更新 8 字段）', bold=True)

add_h2('4.7 测试铁律')
add_bullet('测试前必 docker compose restart app（CLAUDE.md 752 行铁律）', bold=True)
add_bullet('新增 .vue/.css 字面量用 Edit tool 逐文件改，不用 PowerShell 批量 Set-Content（UTF-8 BOM 污染）', bold=True)
add_bullet('新增测试必须有"原始错误样本 + 真实正确样本"双向测试', bold=True)
add_bullet('LLM 测评不能依赖单次（2 轮稳定性测试 0% 一致性是已知问题）', bold=True)


# ============================================================
# 第五部分：待开发的内容
# ============================================================
doc.add_page_break()
add_h1('五、待开发的内容（CLAUDE.md "不在本次范围"汇总）')

add_para('以下内容均已识别但受限于时间/风险/优先级尚未实现，按优先级分 3 类：', bold=True)

# --- 5.1 P0 优先 ---
add_h2('5.1 P0 高优先级（核心架构收尾）')

add_h3('5.1.1 MeetingView 1088 行拆分（CLAUDE.md 多次 defer）')
add_para('用户决策信号：当前 MeetingView 包含 6 个 dialog（MeetingTemplateDialog / MeetingMinutesDialog / MeetingCreateDialog / VoiceprintEnrollDialog / PasteAnalyzeDialog / AgentTracesDialog）+ 复杂 CRUD + 听会状态机，组件过大难以维护。')
add_bullet('预计 2-3h 重构风险高')
add_bullet('拆为 5-6 个子组件：MeetingList / MeetingCreateForm / MeetingLiveBanner / MeetingFilterBar / MeetingActionBar')
add_bullet('预期收益：维护成本降低 + 测试粒度细化 + dark mode 适配更精准')

add_h3('5.1.2 agentic_loop.py 1123 行拆分（后端核心模块）')
add_para('用户决策信号：当前 agentic_loop.py 包含 Phase 0/1/2 + 流式 abort + JSONB flag_modified + Self-RAG + 4 域 fan-out + intent_classifier 集成，单文件 1123 行。')
add_bullet('预计 3-4h 重构（后端核心模块，影响范围大）')
add_bullet('拆为 4-5 个模块：intent_classifier.py / phase0_planning.py / phase1_execution.py / phase2_synthesis.py / abort_handler.py')
add_bullet('预期收益：单测粒度更细 + Phase 0/1/2 独立演进 + 与方案 C 6 铁律更清晰对齐')

add_h3('5.1.3 后端 alembic 033 / agent_traces 清理（后端运维轮次）')
add_para('用户决策信号：当前 alembic 链存在多 head 状态（033_mvh 是 DB stamp 但本地迁移文件丢失，仅 alembic_history 中可见）。')
add_bullet('alembic 033_mvh 文件重建')
add_bullet('agent_traces 30 天清理策略实现（与 reminder / chat_history 对齐）')
add_bullet('预期收益：迁移链稳定 + 自动清理减少运维负担')

# --- 5.2 P1 中优先级 ---
add_h2('5.2 P1 中优先级（功能完善 + 体验优化）')

add_h3('5.2.1 Mobile 6 路由 baseline 扩')
add_para('当前 mobile 只有 dashboard / chat / knowledge / tasks / meetings / settings 6 路由 baseline。')
add_bullet('扩展路由：projects / members / project-stats / member-detail / project-detail / meeting-detail / meeting-room / knowledge-detail 等')
add_bullet('预计 4-6h 收 baseline + Playwright spec')
add_bullet('预期收益：移动端视觉回归覆盖到所有核心页')

add_h3('5.2.2 KnowledgeExtractionsPanel / KnowledgeImageGallery dark 化扩 baseline')
add_para('当前两个组件 dark 化但未加入 Playwright baseline。')
add_bullet('knowledge detail page dark 化扩 baseline（image gallery + extraction panel）')
add_bullet('预计 2h')
add_bullet('预期收益：阅读器完整 dark mode 视觉回归覆盖')

add_h3('5.2.3 T3.5 qa-bench 性能优化（需主循环重构）')
add_para('当前 qa-bench 200 题 smoke 套件 5min 80% 阈值，主循环重构后可优化到 3min 80%。')
add_bullet('重构 agentic_loop 主循环（计划 → 执行 → 总结 → done）')
add_bullet('并行工具调用（当前串行，4 域 fan-out 时 4x 时间）')
add_bullet('预期收益：qa-bench smoke 提速 40%')

add_h3('5.2.4 模板批量复制 builtin 为 custom')
add_para('当前内置 4 个模板（组会/一对一/立项会/自由）用户如需基于 builtin 微调，必须单击 4 次单独复制。')
add_bullet('添加"批量复制为自定义"按钮（TemplatesView）')
add_bullet('预计 1-2h')
add_bullet('预期收益：模板管理 UX 提升')

# --- 5.3 P2 低优先级 ---
add_h2('5.3 P2 低优先级（功能扩展 / 边缘场景）')

add_h3('5.3.1 Web Push / Periodic Background Sync（投资回报低）')
add_para('当前 PWA Background Sync 已实现 Local Notification，Web Push 协议需服务端 + 浏览器双端配置。')
add_bullet('Web Push 协议（VAPID 密钥 + 浏览器订阅）')
add_bullet('Periodic Background Sync（定时后台任务，需要 Chrome 支持）')
add_bullet('预期收益：消息推送可达性（但当前 Local Notification 已够用）')
add_para('plan 决策不投资（投资回报低 / 浏览器支持窄）', italic=True)

add_h3('5.3.2 模板导出/导入')
add_bullet('TemplatesView 添加"导出选中模板为 JSON" + "从 JSON 导入"')
add_bullet('跨团队知识共享场景')

add_h3('5.3.3 模板版本控制')
add_bullet('模板历史版本存储 + 一键回滚')
add_bullet('团队成员误删模板场景')

add_h3('5.3.4 模板审批流')
add_bullet('模板创建/修改需组长/管理员审批')
add_bullet('模板质量保证')

add_h3('5.3.5 模板使用统计')
add_bullet('TemplatesView 添加"最近使用时间" + "使用次数"列')
add_bullet('TemplatesView 添加"热门模板"排序')

add_h3('5.3.6 模板标签/分类')
add_bullet('模板标签系统（自定义标签如"周会/项目评审/技术分享"）')
add_bullet('当前 cloned_from_id 标签够用，但业务场景可扩展')

add_h3('5.3.7 模板按角色可见')
add_bullet('不同角色看到不同模板（普通成员 vs 管理员）')
add_bullet('当前 全部人可见够用，权限控制可未来加')

add_h3('5.3.8 模板 audit log')
add_bullet('模板操作审计日志（谁在什么时候改了什么）')
add_bullet('细粒度操作追溯')

add_h3('5.3.9 模板拖拽排序')
add_bullet('TemplatesView 表格行支持拖拽排序')
add_bullet('overkill，4 个内置 + 10 个自定义，UI 表格已够')

add_h3('5.3.10 移动端独立 MobileTemplatesView')
add_para('当前 /admin/templates 是桌面端 el-table，移动端无独立页面。')
add_bullet('MobileTemplatesView 列表 + 卡片视图切换')
add_bullet('用户决策"只做桌面"，但未来可扩展')

add_h3('5.3.11 移动端 long-press 模板管理（MobileTemplatesView）')
add_bullet('仿照 MobileCreateDialog long-press 模式')
add_bullet('与 v77 P2.6-G.1 同模式扩展')

add_h3('5.3.12 12 处剩余 runtime :style 全量收敛（CSS-in-JS 收尾）')
add_para('当前 plan 估 16 处，P2.6-E.1 收敛 8 处，剩 12 处：priority/status/badge/conf-bar/quick-icon/theme-preview/card-hero 等单步风险高。')
add_bullet('逐文件渐进替换（不用 PowerShell 批量）')
add_bullet('预计 4-6h')

add_h3('5.3.13 136 处缓动字面量全量替换（E.2 收尾）')
add_para('P2.6-E.2 实际替换 70 处（plan 估 145 误算），剩 95 处：transition: all 0.Xs ease / 200ms var(--ease-out) 等含 ease 关键字 / ms 单位场景。')
add_bullet('用 Edit tool 逐文件改（不能用 PowerShell 批量）')
add_bullet('预计 6-8h')

add_h3('5.3.14 跨域用户场景（member_id 维度限流）')
add_para('当前限流按 user_id 维度，但 member_id 维度（每个成员在团队中的行为）未限流。')
add_bullet('添加 member_id 维度限流（防止单一成员刷接口）')
add_bullet('预计 2h')

# --- 5.4 P3 探索性 ---
add_h2('5.4 P3 探索性（产品/技术演进）')

add_h3('5.4.1 #009 Self-RAG 离线重检索决策')
add_para('当前 Self-RAG 仅在对话时实时 judge + 重检索，离线重检索（定时任务对历史对话）未实现。')
add_bullet('Celery 任务：每日扫描昨日对话，识别低质响应')
add_bullet('离线重检索 + 增量更新')
add_bullet('预计 4-6h')

add_h3('5.4.2 #009 Phase 1 持续微调（Haiku judge 优化）')
add_bullet('当前 Haiku judge 800ms 偏慢')
add_bullet('改用规则 + 小模型混合')
add_bullet('预计 3-4h')

add_h3('5.4.3 智能论文阅读器 v28 step 9+（P1 工作）')
add_para('v26/v26.1/v27/v28 收官后，plan 中 P1 工作（智能图文匹配 / 翻译 / 知识图谱 / 段落操作 / AI 总结）按用户要求暂不启动。')
add_bullet('智能图文匹配：anchor + figure 智能配对')
add_bullet('翻译：双语对照')
add_bullet('知识图谱：阅读时显示关联条目')
add_bullet('段落操作：注释 / 摘录 / 引用')
add_bullet('AI 总结：自动摘要 + 关键概念')
add_bullet('预计 10-20h')

add_h3('5.4.4 知识库动态分类体系涌现')
add_para('当前 dynamic_taxonomy_service 已实现，但涌现分类的展示 UI 未做。')
add_bullet('KnowledgeView 分类 tab 添加"动态涌现"列')
add_bullet('预计 4-6h')

add_h3('5.4.5 知识库自主研究引擎扩展')
add_para('当前 auto_research_service 已实现，但搜索源只有搜狗+必应。')
add_bullet('扩展搜索源（Google Scholar / arXiv / Semantic Scholar）')
add_bullet('预计 6-8h')

add_h3('5.4.6 移动端 PWA 4 策略扩展')
add_para('当前 4 策略：manifest + service worker + useSafeArea + IndexedDB')
add_bullet('Background Fetch API（断网下载大文件）')
add_bullet('Web Share API（移动端分享到系统）')
add_bullet('预计 4-6h')

add_h3('5.4.7 LLM 元话语/thinking 增强')
add_para('当前 19 种元话语 pattern + 兜底剥除')
add_bullet('扩展到 30+ 种（基于 qa-bench 失败题分析）')
add_bullet('动态学习（基于负反馈优化）')
add_bullet('预计 3-4h')

add_h3('5.4.8 深度学习方向 125+ 项（6 个月路线图）')
add_para('A/B/C 三类 + 智能体 + 备用资产 + GPU 调度，PyTorch > TF，跳过 MATLAB')
add_bullet('A 类（基础能力）40+ 项')
add_bullet('B 类（高级能力）35+ 项')
add_bullet('C 类（智能体）35+ 项')
add_bullet('预计 6 个月 + 多 GPU')

add_h3('5.4.9 #043 进一步扩展')
add_para('当前 8 phase 收官（基础 + 旧数据迁移 + UI 升级 + 清理）')
add_bullet('P0: 跨设备 WebSocket push（替代 MVP 拉取）')
add_bullet('P1: 消息全文搜索（PG tsvector + GIN 索引）')
add_bullet('P2: 多模态消息（图片/语音/文件）')
add_bullet('预计 8-12h')


# ============================================================
# 第六部分：申请更大 Token Plan 计划的理由
# ============================================================
doc.add_page_break()
add_h1('六、申请更大 Token Plan 计划的理由')

add_h2('6.1 项目规模需要')
add_para('当前 313,000+ 行代码 / 799 文件 / 1,545+ commits / 46 开发天数，CLAUDE.md 沉淀 566KB（50+ 章节铁律），75+ 沉淀记忆。')
add_para('单次会话常因以下场景触发 token 限制：', bold=True)
add_bullet('读取 CLAUDE.md 全量（566KB ≈ 14.5 万 token）')
add_bullet('读取 CHANGELOG.md（92KB ≈ 2.4 万 token）')
add_bullet('读取多文件上下文（app/agent/ 20+ 模块 + scripts/ 25+ 脚本）')
add_bullet('读取 50+ memory 沉淀（每个 5-20KB）')
add_bullet('输出 4 子任务收官报告（每次 2-3 万 token）')
add_bullet('处理 1545+ commits 的 git 历史回溯')
add_para('在大型多 PR 收官时，常因单次会话输出超限被截断，**必须分多次会话接力**，影响效率。', bold=True)

add_h2('6.2 任务复杂度需要')
add_para('项目主线任务常需 2-3 万 token 输出，包含：', bold=True)
add_bullet('4 子任务代码改动（每个 800-2000 行）')
add_bullet('8-10 个文件改动 + 端到端验证')
add_bullet('5-15 条新铁律沉淀')
add_bullet('Memory 文件创建（每个 5-15KB）')
add_bullet('commit 链梳理 + 状态报告')
add_bullet('CLAUDE.md 章节更新（每次 200-500 行）')
add_para('当前 token 限制下，单次完整收官一个 P2.6 子任务需要 2-3 轮精打细算，影响质量。', bold=True)

add_h2('6.3 开发效率需要')
add_para('当前 46 天内已完成 1545+ commits，平均 33 commits/天。')
add_para('高峰期（如 v77 P2.6-G 系列 6 子任务收官、#043 8 phase 收官）单日 20+ commits + 10+ memory 沉淀 + 多次精打细算。')
add_para('如能获得更大 token plan：', bold=True)
add_bullet('单次会话可完成完整 4 子任务收官（避免 2-3 轮精打细算）')
add_bullet('可读取更多上下文（CLAUDE.md + CHANGELOG.md + memory 全部）而无需选择性读取')
add_bullet('可输出更完整的报告（铁律 + memory + 5 维度验证 + 踩坑教训）')
add_bullet('可减少精打细算导致的"漏记铁律"风险')
add_bullet('可支持更多端到端验证（Playwright + Vitest + pytest）')

add_h2('6.4 长期演进需要')
add_para('项目未来 6 个月规划：', bold=True)
add_bullet('深度学习方向 125+ 项（6 个月路线图）')
add_bullet('#009 Self-RAG 持续微调（Phase 0.5 + 1 + 离线重检索）')
add_bullet('#043 进一步扩展（WebSocket push + 全文搜索 + 多模态消息）')
add_bullet('智能论文阅读器 v28 P1 工作（智能图文匹配 / 翻译 / 知识图谱）')
add_bullet('Mobile 路由 baseline 扩（10+ 路由）')
add_bullet('后端限流 v31.3 系列（动态配置 + 灰度）')
add_bullet('视觉/代码质量收官（MeetingView / agentic_loop 拆分）')
add_para('总预估：200-300h 工作量，50+ 收官报告，200+ 铁律沉淀。', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))

add_h2('6.5 当前 token 使用痛点')
add_bullet('单次精打细算只能写 4-6 个文件 + 1-2 个测试，4 子任务收官需 2-3 轮接力')
add_bullet('CLAUDE.md 全量读取经常触发 token 限制，需分次读取影响上下文连贯性')
add_bullet('memory 文件批量沉淀时需要精简描述，影响后续检索质量')
add_bullet('Playwright / Vitest 输出长日志时需要截断，影响问题定位精度')
add_bullet('多 PR 收官时铁律沉淀常因 token 不足被简化，未来再补全成本高')
add_bullet('git log + git diff 输出长时需要分页，影响 commit 链分析')

add_h2('6.6 期望的 token plan 改进')
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
add_table_header(table, ['当前', '期望', '收益'])
improvements = [
    ('单次输出 8K tokens', '单次输出 16-32K tokens', '完整 4 子任务收官报告'),
    ('CLAUDE.md 选择性读', 'CLAUDE.md 全量读（566KB）', '上下文连贯 + 铁律完整'),
    ('CHANGELOG 分次读', 'CHANGELOG 全量读（92KB）', '历史回溯 + commit 链分析'),
    ('memory 选择性读', 'memory 批量读（75+ 条）', '沉淀质量提升 + 复用效率'),
    ('Playwright 输出截断', 'Playwright 全量日志', '问题定位精度 + e2e 调试'),
    ('4 子任务 2-3 轮接力', '4 子任务 1 轮收官', '开发效率 +50%'),
    ('git log 分页', 'git log 全量', 'commit 链分析 + 影响面评估'),
]
for cur, exp, benefit in improvements:
    add_table_row(table, [cur, exp, benefit])

doc.add_paragraph()

add_h2('6.7 申请结论')
add_para('基于以上 6 个维度分析，建议将 token plan 升级为：', bold=True, size=12)
add_bullet('单次输出提升到 16-32K tokens（当前 8K）', bold=True)
add_bullet('上下文窗口提升到 400-500K tokens（当前 200K）', bold=True)
add_bullet('保留所有现有功能（CLAUDE.md 读取、memory 沉淀、Playwright 调试等）', bold=True)
add_para('预期收益：', bold=True, size=11)
add_bullet('开发效率提升 50%（4 子任务 1 轮 vs 2-3 轮）', bold=True)
add_bullet('沉淀质量提升（铁律完整 + memory 全量）', bold=True)
add_bullet('调试效率提升（Playwright 全量日志 + git log 全量）', bold=True)
add_bullet('6 个月规划可顺利推进（200-300h 工作量）', bold=True)


# ============================================================
# 第七部分：附录
# ============================================================
doc.add_page_break()
add_h1('七、附录')

add_h2('7.1 关键文档索引')
add_bullet('CLAUDE.md — 项目上下文（566KB，50+ 章节铁律）', bold=True)
add_bullet('ROADMAP.md — 路线图')
add_bullet('README.md — 项目说明')
add_bullet('CHANGELOG.md — 变更日志（92KB）')
add_bullet('docs/ — 设计文档（21 篇）')
add_bullet('memory/ — 75+ 沉淀记忆')
add_bullet('plans/ — 10+ 实现计划')
add_bullet('scripts/ — 25+ 工具脚本')

add_h2('7.2 关键 scripts 脚本清单')
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
add_table_header(table, ['脚本', '用途'])
scripts = [
    ('scripts/migrate_kb_tags.py', 'KB 数据清洁：自动生成内容 tags 归并 + 测试/TEST 样板删除（303 行）'),
    ('scripts/migrate_kb_dedup_titles.py', 'KB 物理删字节相同副本（FK 防御 + JSON 备份 ~560 行）'),
    ('scripts/migrate_kb_source_type.py', '180 张 [拓展-XX] 卡片 source_type 重分类（226 行）'),
    ('scripts/auto_intake_rollback.py', '7 天自动清理 source_type=auto_expansion（KB 入库 rollback）'),
    ('scripts/reprocess_meeting.py', '会议发言人重处理 9 步 CLI（ERes2Net + KMeans + file 备份）'),
    ('scripts/verify_login_redis.py', 'login_limiter Redis 化 + Retry-After 响应头验证'),
    ('scripts/verify_v31_2_5_restart.py', '抗 docker restart 验证'),
    ('scripts/verify_cross_meeting_recognition.py', '声纹 strict merge 跨会议识别率 90% 门禁验证'),
    ('scripts/restore_voiceprint.py', '撤销单条 entry 但保留更早 merges'),
    ('scripts/verify_v28_figures.py', 'v28 论文图片结构化字段 5 大验证维度'),
    ('scripts/benchmark_asr.py', 'ASR 模型对比 benchmark（Whisper vs SenseVoice vs Paraformer）'),
    ('scripts/check-token-orphans.sh', 'CSS token orphan 检测（--ci-mode 输出 GH Actions annotation）'),
    ('scripts/check-dist-before-commit.sh', 'pre-commit hook auto-add web/dist/'),
    ('scripts/replace-easing-literals.js', 'Node.js 缓动字面量 token 化脚本（CLAUDE.md PowerShell UTF-8 BOM 第 4 次教训强化）'),
    ('scripts/replace-transition-all-literals.js', 'Node.js transition: all token 化脚本'),
    ('scripts/deploy-auto.sh', '云服务器 webhook 自动部署脚本（带 .env.webhook 守卫）'),
    ('scripts/auto_intake_rollback.py', 'KB 入库 rollback（7 天自动清理）'),
    ('scripts/qa-bench/runner.py', 'qa-bench 跑测主入口（7 维评分 + 3 P0 检测器）'),
    ('scripts/qa-bench/onebyone.py', 'qa-bench 逐个问答'),
    ('scripts/qa-bench/gen500.py', 'qa-bench 动态生成 500 题'),
    ('scripts/qa-bench/save_to_kb.py', 'qa-bench 答题后自动入库（5 道防线）'),
    ('scripts/qa-bench/view.py', 'qa-bench 答题查看'),
    ('scripts/gen_advanced_report.py', 'qa-bench 高级能力报告生成'),
]
for name, purpose in scripts:
    add_table_row(table, [name, purpose])

add_h2('7.3 关键 alembic 迁移清单')
add_bullet('alembic/versions/019_reminder_ack_snooze_v2.py — 任务提醒体系 v2（6 列）')
add_bullet('alembic/versions/020_knowledge_multimodal.py — Phase 7 多模态知识库（2 表）')
add_bullet('alembic/versions/028_figure_structured_fields.py — v28 论文图片结构化字段（12 列 + 2 索引）')
add_bullet('alembic/versions/033_mvh.py — DB stamp（多 head 警告）')
add_bullet('alembic/versions/034_reset_voice_sample_count.py — 声纹 sample_count 重置为 1')
add_bullet('alembic/versions/038_tpl_cloned_from.py — 模板 cloned_from_id 字段（v77 P2.6-F.5）')
add_bullet('alembic/versions/039_chat_history.py — #043 账号持久化聊天历史（chat_sessions / chat_messages / chat_shares 三表）')

add_h2('7.4 关键 Playwright 测试')
add_bullet('tests/visual/desktop/visual-regression.spec.mjs — Desktop 9 路由 baseline（已废弃 CI 跑，本地 dev 保留）')
add_bullet('tests/visual/mobile/visual-regression.spec.mjs — Mobile 6 路由 baseline')
add_bullet('tests/visual/desktop/v77-p2-6-f-2-regression.spec.mjs — v77 P2.6-F.2 16 项 smoke test')
add_bullet('tests/visual/desktop/v77-p2-6-g-2-templates.spec.mjs — 模板管理 B-17~B-20')
add_bullet('tests/visual/mobile/mobile-long-press.spec.mjs — 移动端 long-press M-13~M-16')

add_h2('7.5 关键 Vitest 测试套件')
add_bullet('web/src/views/MeetingView/__tests__/ — 19 个测试（F.3 8 + F.4 4 + F.5 4 + 旧 3）')
add_bullet('web/src/views/admin/__tests__/TemplatesView.test.js — 17 PASSED')
add_bullet('web/src/components/chat/blocks/__tests__/HypothesisBlock.spec.js — 14 PASSED（v76.3 组件级 CSS variable 测试）')
add_bullet('web/src/utils/__tests__/paperAdapter.test.js — 73+10 = 83 PASSED（v26/v27/v28 系列）')
add_bullet('web/src/composables/chat/__tests__/useChatStream.test.js — 多会话并行 8 铁律')
add_bullet('web/src/composables/__tests__/useChatMigration.test.js — 9 PASSED（#043 旧数据迁移）')
add_bullet('web/src/composables/__tests__/useGlobalShortcuts.test.js — 9 PASSED（#043 全局快捷键）')
add_bullet('web/src/composables/__tests__/useKbMonitor.test.js — KB 入库监控')

add_h2('7.6 关键 pytest 测试套件')
add_bullet('tests/test_auth.py — 8 PASS（conftest.py 隔离 bug 修复后）')
add_bullet('tests/test_meeting_template_service.py — 9 PASS（v77 P2.6-F.5）')
add_bullet('tests/test_migrate_kb_tags.py — 14 PASS（KB 数据清洁）')
add_bullet('tests/test_migrate_kb_dedup_titles.py — 19 PASS（KB 物理删字节相同副本）')
add_bullet('tests/test_migrate_kb_source_type.py — 7 PASS（180 张 [拓展-XX] 卡片重分类）')
add_bullet('tests/test_chat_history_service.py — 24 test PASS（#043 Phase 8）')
add_bullet('tests/test_chat_history_tasks.py — 7 test PASS（#043 Phase 7 Celery 30 天清理）')


# ============================================================
# 保存
# ============================================================
output_path = '/e/microbubble-agent/docs/MicroBubble_Agent_Dev_Status_Report_2026-06-30.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
# Windows GBK 编码安全输出
import sys
sys.stdout.reconfigure(encoding='utf-8')
print(f"[OK] Word 文档生成成功: {output_path}")
print(f"     文件大小: {os.path.getsize(output_path) / 1024:.1f} KB")
